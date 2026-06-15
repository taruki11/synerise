from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def render_markdown_to_docx(markdown_path: Path, docx_path: Path) -> None:
    lines = markdown_path.read_text(encoding='utf-8').splitlines()
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    in_code_block = False
    for line in lines:
        stripped = line.rstrip('\n')

        if stripped.startswith('```'):
            in_code_block = not in_code_block
            if not in_code_block:
                document.add_paragraph('')
            continue

        if in_code_block:
            p = document.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            continue

        if not stripped.strip():
            document.add_paragraph('')
            continue

        if stripped.startswith('# '):
            p = document.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(stripped[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
            continue

        if stripped.startswith('## '):
            p = document.add_paragraph()
            run = p.add_run(stripped[3:].strip())
            run.bold = True
            run.font.size = Pt(13)
            continue

        if stripped.startswith('### '):
            p = document.add_paragraph()
            run = p.add_run(stripped[4:].strip())
            run.bold = True
            run.font.size = Pt(11.5)
            continue

        if stripped.startswith('- '):
            p = document.add_paragraph(style='List Bullet')
            p.add_run(stripped[2:].strip())
            continue

        p = document.add_paragraph()
        p.add_run(stripped)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit('Usage: python scripts/export_report_docx.py <input.md> <output.docx>')
    render_markdown_to_docx(Path(sys.argv[1]), Path(sys.argv[2]))


if __name__ == '__main__':
    main()
